from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any


# ------------------------------------------------------------
# STYLE HELPERS
# ------------------------------------------------------------

def _add_title(doc, title: str) -> None:
    """Centered, blue, bold title."""
    p = doc.add_heading(title, level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color


def _add_section_title(doc, title: str) -> None:
    """Professional section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)  # Adjust before space
    p.paragraph_format.space_after = Pt(6)   # Adjust after space
    run = p.add_run(title)
    run.font.size = Pt(14)
    run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_paragraph(doc, text: str) -> None:
    """Left-aligned clean paragraph."""
    p = doc.add_paragraph(text)
    fmt = p.paragraph_format
    fmt.space_after = Pt(6)  # More space after each paragraph
    fmt.space_before = Pt(1)  # Small space before
    fmt.line_spacing = 1.15  # Adjusted for better readability
    fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align text


def _add_bullet(doc, text: str) -> None:
    """Proper Word bullet point (no asterisks)."""
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = Pt(11)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align bullets


# ------------------------------------------------------------
# EXPORT MAIN FUNCTION
# ------------------------------------------------------------

def save_docx_report(save_path: str, sections: List[Dict[str, Any]], subset: str) -> str:
    """Creates a polished DOCX diagnostic report based on structured LLM output."""

    doc = Document()

    # Clean 1-inch margins
    for sec in doc.sections:
        sec.top_margin = Inches(0.8)
        sec.bottom_margin = Inches(0.8)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Add title once with blue color
    _add_title(doc, f"LLM Engine Health Diagnostic Report — {subset}")

    for sec in sections:
        # Section Title
        _add_section_title(doc, sec["title"])

        # Section paragraph
        if sec.get("content"):
            _add_paragraph(doc, sec["content"])

        # Bullet points (sensor deviations)
        if sec.get("bullet_points"):
            for bullet in sec["bullet_points"]:
                _add_bullet(doc, bullet)

    doc.save(save_path)
    return save_path

# ============================================================================